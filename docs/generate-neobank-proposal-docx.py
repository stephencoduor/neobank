#!/usr/bin/env python3
"""Generate NeoBank-Digital-Banking-Proposal.docx using python-docx."""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_PATH = os.path.join(OUTPUT_DIR, "NeoBank-Digital-Banking-Proposal.docx")

# Savanna brand colors
PRIMARY = RGBColor(0x2D, 0x6A, 0x4F)
PRIMARY_DARK = RGBColor(0x1B, 0x43, 0x32)
GOLD = RGBColor(0xE9, 0xB9, 0x49)
LIGHT_GREEN = RGBColor(0xD8, 0xF3, 0xDC)
TEXT_HEADING = RGBColor(0x1A, 0x1A, 0x1A)
TEXT_BODY = RGBColor(0x37, 0x41, 0x51)
TEXT_MUTED = RGBColor(0x6B, 0x72, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DANGER = RGBColor(0xEF, 0x44, 0x44)
WARNING = RGBColor(0xEA, 0xB3, 0x08)
SUCCESS = RGBColor(0x22, 0xC5, 0x5E)


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = PRIMARY if level <= 2 else PRIMARY_DARK
        run.font.name = "Calibri"
    return h


def add_body(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = TEXT_BODY
    run.bold = bold
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        run.font.color.rgb = TEXT_BODY
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = TEXT_BODY
    return p


def add_gold_line(doc):
    p = doc.add_paragraph()
    run = p.add_run("_" * 30)
    run.font.color.rgb = GOLD
    run.font.size = Pt(6)
    p.paragraph_format.space_after = Pt(8)


def styled_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        run.font.color.rgb = WHITE
        set_cell_shading(cell, "2D6A4F")

    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = "Calibri"
            run.font.color.rgb = TEXT_BODY
            if r % 2 == 1:
                set_cell_shading(cell, "FAFAF5")

    return table


def gap_label(level):
    return level


def build_docx():
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = TEXT_BODY

    # ──── COVER PAGE ──────────────────────────────────────────────────────────
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Qsoftwares Ltd")
    run.font.size = Pt(14)
    run.font.color.rgb = PRIMARY
    run.bold = True
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    run = p.add_run("Digital Financial Solutions")
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT_MUTED
    run.font.name = "Calibri"

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Next-Gen Digital Banking\n& Payments Ecosystem")
    run.font.size = Pt(28)
    run.font.color.rgb = PRIMARY_DARK
    run.bold = True
    run.font.name = "Calibri"

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Technical Proposal  |  Gap Analysis  |  Implementation Plan")
    run.font.size = Pt(13)
    run.font.color.rgb = PRIMARY
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    run = p.add_run("Mobile-First Financial Operating System for Emerging Markets")
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT_MUTED
    run.italic = True
    run.font.name = "Calibri"

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("April 2026  |  20-Week Delivery  |  $60,000 Fixed Price")
    run.font.size = Pt(12)
    run.font.color.rgb = GOLD
    run.bold = True
    run.font.name = "Calibri"

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("CONFIDENTIAL")
    run.font.size = Pt(10)
    run.font.color.rgb = DANGER
    run.bold = True
    run.font.name = "Calibri"

    doc.add_page_break()

    # ──── TABLE OF CONTENTS ───────────────────────────────────────────────────
    add_heading_styled(doc, "Table of Contents", level=1)
    add_gold_line(doc)
    toc = [
        "1. Executive Summary",
        "2. Platform Readiness Assessment",
        "3. Gap Analysis \u2014 Requirements Matrix",
        "4. Critical Gaps & Solutions",
        "5. BaaS & Sponsor Bank Strategy",
        "6. Recommended Technical Stack",
        "7. Implementation Phases (20 Weeks)",
        "8. Budget Breakdown",
        "9. Third-Party Dependencies",
        "10. Risk Assessment",
        "11. Our Competitive Advantage",
        "12. Edge Case Matrix",
        "13. Deliverables & Next Steps",
    ]
    for item in toc:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(12)
        run.font.name = "Calibri"
        run.font.color.rgb = TEXT_BODY
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ──── SECTION 1: Executive Summary ────────────────────────────────────────
    add_heading_styled(doc, "1. Executive Summary", level=1)
    add_gold_line(doc)
    add_body(doc,
        "We propose building a comprehensive digital banking and payments ecosystem "
        "for an emerging market, leveraging our existing Apache Fineract core banking "
        "infrastructure, 9 payment provider integrations, and proven mobile banking "
        "patterns. Our platform already covers ~60% of the required backend "
        "infrastructure, dramatically reducing build time from 12+ months to 20 weeks.")

    add_body(doc,
        "Key Decision: We recommend Flutter for the cross-platform mobile app "
        "(iOS + Android from a single codebase), with Riverpod state management "
        "and Drift for offline-first local storage.", bold=True)

    add_body(doc,
        "The platform serves as a high-utility retail banking app for individuals "
        "and a robust payment-acceptance suite for merchants, replacing cash-heavy "
        "legacy systems with a secure, mobile-first Financial Operating System.")

    doc.add_paragraph()
    styled_table(doc, ["Metric", "Value"], [
        ["Backend Infrastructure Ready", "~60%"],
        ["Payment Providers Integrated", "9 (Africa + Asia + Global)"],
        ["Delivery Timeline", "20 weeks (5 months)"],
        ["Fixed Budget", "$60,000"],
    ])

    doc.add_page_break()

    # ──── SECTION 2: Platform Readiness ───────────────────────────────────────
    add_heading_styled(doc, "2. Platform Readiness Assessment", level=1)
    add_gold_line(doc)
    add_body(doc,
        "Our existing platform comprises three production applications with a shared "
        "Apache Fineract backend. These capabilities transfer directly to the NeoBank project:")

    styled_table(doc, ["Capability", "Details", "Status"], [
        ["Core Banking Engine", "Fineract \u2014 double-entry GL, loans, savings, multi-currency", "Production"],
        ["Payment Integrations", "M-Pesa, Airtel, MTN, Flutterwave, Paystack, Cellulant, AT, Razorpay, Stripe", "Production"],
        ["React Admin Dashboard", "Client CRUD, loan/savings management, mobile money, reports", "Production"],
        ["Tax Compliance", "eTIMS/KRA integration \u2014 invoices, credit notes, tax codes", "Production"],
        ["Credit Scoring", "PataScore \u2014 bank statement + M-Pesa analysis", "Production"],
        ["Client Management", "KYC document storage, client lifecycle, group/center hierarchy", "Production"],
        ["Mobile UX Patterns", "50+ screens: auth, dashboard, loans, payments, QR, biometric", "Proven"],
        ["Multi-Tenant Architecture", "Isolate data per organization from day one", "Production"],
    ])

    doc.add_page_break()

    # ──── SECTION 3: Gap Analysis ─────────────────────────────────────────────
    add_heading_styled(doc, "3. Gap Analysis \u2014 Requirements Matrix", level=1)
    add_gold_line(doc)

    add_heading_styled(doc, "A. Digital Banking Infrastructure", level=2)
    styled_table(doc, ["Requirement", "Current State", "Gap Level"], [
        ["Tiered Accounts (Personal + Business)", "Fineract supports client types + savings products", "LOW"],
        ["Automated KYC/AML (ID + liveness)", "Document storage only, no automated verification", "HIGH"],
        ["Shadow Ledgering (real-time)", "Full GL exists; pending transaction layer needed", "MEDIUM"],
        ["Multi-currency reconciliation", "Multi-currency supported; FX feeds needed", "LOW"],
    ])
    doc.add_paragraph()

    add_heading_styled(doc, "B. Card Management System", level=2)
    styled_table(doc, ["Requirement", "Current State", "Gap Level"], [
        ["Physical prepaid debit cards", "No card infrastructure", "CRITICAL"],
        ["Virtual instant cards", "None", "CRITICAL"],
        ["Freeze/unfreeze, spend limits, alerts", "None", "CRITICAL"],
        ["Contactless (NFC) + Chip & PIN", "None \u2014 requires card issuing partner", "CRITICAL"],
    ])
    doc.add_paragraph()

    add_heading_styled(doc, "C. Merchant & POS Solutions", level=2)
    styled_table(doc, ["Requirement", "Current State", "Gap Level"], [
        ["Bluetooth POS terminals", "No POS infrastructure", "HIGH"],
        ["SoftPOS / Tap-to-Phone", "None", "HIGH"],
        ["Instant settlement", "Mobile money B2C disbursement exists", "MEDIUM"],
    ])
    doc.add_paragraph()

    add_heading_styled(doc, "D. P2P Ecosystem", level=2)
    styled_table(doc, ["Requirement", "Current State", "Gap Level"], [
        ["Instant P2P via phone/alias/QR", "QR screen + basic transfers exist", "MEDIUM"],
        ["QR for P2P + Merchant payments", "Basic QR in mobile app", "MEDIUM"],
    ])
    doc.add_paragraph()

    add_heading_styled(doc, "E. Technical & Security Standards", level=2)
    styled_table(doc, ["Requirement", "Current State", "Gap Level"], [
        ["SOC2/PCI-DSS compliance", "No compliance framework", "CRITICAL"],
        ["OAuth2 + JWT authentication", "Basic auth only", "HIGH"],
        ["Data residency compliance", "Not architected for regional deployment", "HIGH"],
        ["Mobile security (cert pinning)", "Biometric exists; cert pinning missing", "MEDIUM"],
        ["High-concurrency backend", "Fineract Spring Boot + modular arch", "LOW"],
    ])

    doc.add_page_break()

    # ──── SECTION 4: Critical Gaps & Solutions ────────────────────────────────
    add_heading_styled(doc, "4. Critical Gaps & Solutions", level=1)
    add_gold_line(doc)

    add_heading_styled(doc, "4.1 Card Issuing & Management", level=2)
    add_body(doc,
        "This is the most significant gap. Card infrastructure requires a BaaS partner "
        "already PCI-DSS certified and connected to card networks.")
    styled_table(doc, ["Region", "Card Issuing Partners"], [
        ["Africa", "Union54, Flutterwave (virtual cards), Paystack Issuing"],
        ["Asia", "Stripe Issuing, Rapyd, Nium"],
        ["Global", "Marqeta, Stripe Issuing, Adyen Issuing"],
    ])
    add_body(doc,
        "Architecture: Flutter App \u2192 Our Backend \u2192 Card Issuing API \u2192 "
        "Webhook receiver \u2192 Fineract GL posting. Our systems never touch raw card "
        "numbers (tokenization). Effort: 4\u20135 weeks.")

    add_heading_styled(doc, "4.2 KYC/AML Automated Workflows", level=2)
    add_body(doc,
        "Automated identity verification is mandatory. We integrate with Smile Identity "
        "(Africa) or Onfido (Global) for ID verification, liveness detection, OCR, and "
        "AML/sanctions screening.")
    add_body(doc,
        "KYC Flow: Capture ID (front+back) \u2192 Liveness check \u2192 OCR extraction "
        "\u2192 AML screening \u2192 Risk score \u2192 Auto-approve or manual queue "
        "\u2192 Fineract account activation. Effort: 2\u20133 weeks.")

    add_heading_styled(doc, "4.3 PCI-DSS & SOC2 Compliance", level=2)
    add_body(doc,
        "By using a BaaS partner for card issuing, we reduce PCI-DSS scope to SAQ-A. "
        "SOC2 requires encryption at rest, access logging, vulnerability scanning, and "
        "documented security policies. Effort: 3\u20134 weeks.")

    add_heading_styled(doc, "4.4 Merchant & POS Solutions", level=2)
    add_bullet(doc, " PAX A920 / Sunmi V2 Pro \u2014 Bluetooth + 4G terminals", "Hardware POS:")
    add_bullet(doc, " Mastercard Tap on Phone SDK or Visa Tap to Phone", "SoftPOS:")
    add_bullet(doc, " Extend existing M-Pesa/Airtel B2C for instant payouts", "Settlement:")
    add_bullet(doc, " EMVCo QR specification for interoperable payments", "QR Standard:")
    add_body(doc, "Effort: 4\u20135 weeks")

    doc.add_page_break()

    # ──── SECTION 5: BaaS Strategy ───────────────────────────────────────────
    add_heading_styled(doc, "5. BaaS & Sponsor Bank Strategy", level=1)
    add_gold_line(doc)
    add_body(doc,
        "The product requires a licensed sponsor bank that holds deposits, provides "
        "banking rails, issues cards, and ensures regulatory compliance.")

    styled_table(doc, ["Region", "BaaS Partner", "Capabilities", "Card Issuing"], [
        ["East Africa", "Cellulant, Stanbic API", "Deposits, payments", "Via Visa/MC partner"],
        ["West Africa", "Flutterwave, Paystack", "Deposits, cards", "Flutterwave virtual"],
        ["Southern Africa", "Stitch, Investec API", "Open banking", "Investec issuing"],
        ["South Asia", "Razorpay X, Setu", "Neo-banking, UPI", "Partner bank cards"],
        ["Southeast Asia", "Brankas, Rapyd", "Wallets", "Rapyd issuing"],
        ["Global", "Marqeta, Stripe Treasury", "Full stack", "Marqeta/Stripe cards"],
    ])

    doc.add_page_break()

    # ──── SECTION 6: Technical Stack ──────────────────────────────────────────
    add_heading_styled(doc, "6. Recommended Technical Stack", level=1)
    add_gold_line(doc)

    add_heading_styled(doc, "Backend (Extend Existing)", level=2)
    styled_table(doc, ["Component", "Technology", "Rationale"], [
        ["Core Banking", "Apache Fineract (Java 21)", "Already built, proven"],
        ["API Gateway", "Kong / AWS API Gateway", "Rate limiting, auth, routing"],
        ["Auth Server", "Keycloak", "OAuth2 + JWT + MFA"],
        ["Cache", "Redis", "Shadow ledger, sessions"],
        ["Message Queue", "Apache Kafka", "Event sourcing, webhooks"],
        ["Database", "PostgreSQL + Redis", "Fineract-compatible"],
        ["File Storage", "AWS S3", "KYC docs, statements"],
        ["Monitoring", "Grafana + Prometheus", "Full observability"],
    ])
    doc.add_paragraph()

    add_heading_styled(doc, "Flutter Mobile App", level=2)
    styled_table(doc, ["Component", "Package", "Purpose"], [
        ["Framework", "Flutter 3.x (Dart)", "Cross-platform"],
        ["State Mgmt", "Riverpod 2.x", "Compile-safe, testable"],
        ["HTTP Client", "Dio + Retrofit", "Type-safe API calls"],
        ["Local DB", "Drift (SQLite)", "Offline-first"],
        ["Secure Storage", "flutter_secure_storage", "Keychain/Keystore"],
        ["Biometrics", "local_auth", "Fingerprint + Face"],
        ["NFC", "nfc_manager", "Tap-to-pay"],
        ["Camera/OCR", "camera + google_mlkit", "KYC capture"],
        ["QR", "qr_flutter + mobile_scanner", "Generate + scan"],
    ])

    doc.add_page_break()

    # ──── SECTION 7: Implementation Phases ────────────────────────────────────
    add_heading_styled(doc, "7. Implementation Phases (20 Weeks)", level=1)
    add_gold_line(doc)

    add_heading_styled(doc, "Phase 1: Foundation & Security (Weeks 1\u20134)", level=2)
    styled_table(doc, ["Timeline", "Deliverable", "Details"], [
        ["Week 1", "OAuth2 + JWT auth server", "Replace basic auth with Keycloak"],
        ["Week 1\u20132", "Infrastructure hardening", "TLS 1.3, WAF, encryption, VPC"],
        ["Week 2\u20133", "KYC/AML integration", "Smile ID \u2014 ID verification + liveness"],
        ["Week 3\u20134", "BaaS partner vetting", "Evaluate 3\u20135 partners"],
        ["Week 4", "Security documentation", "SOC2 evidence, PCI-DSS SAQ"],
    ])
    doc.add_paragraph()

    add_heading_styled(doc, "Phase 2: Core Banking Enhancement (Weeks 5\u20138)", level=2)
    styled_table(doc, ["Timeline", "Deliverable", "Details"], [
        ["Week 5", "Tiered accounts", "Basic/Standard/Premium tiers"],
        ["Week 5\u20136", "Shadow ledger (Redis)", "Real-time pending balance"],
        ["Week 6\u20137", "P2P payment engine", "Phone lookup, alias, QR"],
        ["Week 7\u20138", "Multi-currency", "FX rates, conversion"],
        ["Week 8", "Notification engine", "FCM + WebSocket alerts"],
    ])
    doc.add_paragraph()

    add_heading_styled(doc, "Phase 3: Card Issuing (Weeks 9\u201312)", level=2)
    styled_table(doc, ["Timeline", "Deliverable", "Details"], [
        ["Week 9\u201310", "Card issuing API", "Virtual + physical cards"],
        ["Week 10\u201311", "Card management UI", "Freeze, limits, PIN"],
        ["Week 11\u201312", "Card transactions", "Auth webhooks, GL posting"],
        ["Week 12", "Card security", "Dynamic CVV, 3DS, tokenization"],
    ])
    doc.add_paragraph()

    add_heading_styled(doc, "Phase 4: Flutter App (Weeks 9\u201316) \u2014 Parallel", level=2)
    styled_table(doc, ["Timeline", "Deliverable", "Details"], [
        ["Week 9\u201310", "App shell + auth", "Flutter, Riverpod, KYC flow"],
        ["Week 11\u201312", "Banking screens", "Dashboard, accounts, P2P"],
        ["Week 13\u201314", "Card screens", "Card UI, NFC"],
        ["Week 14\u201315", "Payment screens", "9 providers, QR"],
        ["Week 15\u201316", "Security + polish", "Cert pinning, testing"],
    ])
    doc.add_paragraph()

    add_heading_styled(doc, "Phase 5: Merchant & POS (Weeks 13\u201316)", level=2)
    styled_table(doc, ["Timeline", "Deliverable", "Details"], [
        ["Week 13", "Merchant onboarding", "KYC, verification, setup"],
        ["Week 13\u201314", "POS integration", "PAX/Sunmi SDK"],
        ["Week 14\u201315", "SoftPOS", "NFC payment acceptance"],
        ["Week 15\u201316", "Settlement engine", "Instant/T+1, MDR fees"],
    ])
    doc.add_paragraph()

    add_heading_styled(doc, "Phase 6: QA & Launch (Weeks 17\u201320)", level=2)
    styled_table(doc, ["Timeline", "Deliverable", "Details"], [
        ["Week 17\u201318", "Automated testing", "Unit + E2E tests"],
        ["Week 18\u201319", "Penetration testing", "Third-party audit"],
        ["Week 19\u201320", "Compliance docs", "SOC2, PCI-DSS, edge cases"],
        ["Week 20", "Soft launch", "Beta release"],
    ])

    doc.add_page_break()

    # ──── SECTION 8: Budget ───────────────────────────────────────────────────
    add_heading_styled(doc, "8. Budget Breakdown", level=1)
    add_gold_line(doc)

    styled_table(doc, ["Phase", "Duration", "Key Deliverables", "Cost"], [
        ["Phase 1: Foundation & Security", "4 weeks", "OAuth2, KYC, infra, BaaS vetting", "$12,000"],
        ["Phase 2: Core Banking", "4 weeks", "Tiered accounts, shadow ledger, P2P, FX", "$12,000"],
        ["Phase 3: Card Issuing", "4 weeks", "Card API, management UI, transactions", "$12,000"],
        ["Phase 4: Flutter App", "8 weeks", "Cross-platform app, 50+ screens", "$12,000"],
        ["Phase 5: Merchant & POS", "4 weeks", "POS, SoftPOS, settlement", "$8,000"],
        ["Phase 6: QA & Launch", "4 weeks", "Testing, pen test, compliance", "$4,000"],
        ["TOTAL", "20 weeks", "Complete Digital Banking Platform", "$60,000"],
    ])

    add_body(doc, "Note: Phases 3 and 4 run in parallel, so total calendar time is 20 weeks (5 months).", bold=True)

    doc.add_page_break()

    # ──── SECTION 9-10: Third-Party + Risk ────────────────────────────────────
    add_heading_styled(doc, "9. Third-Party Dependencies", level=1)
    add_gold_line(doc)
    add_body(doc, "Ongoing operational costs (not included in $60K build cost):")

    styled_table(doc, ["Service", "Provider", "Est. Monthly Cost", "Notes"], [
        ["KYC/AML", "Smile ID / Onfido", "$200\u2013500/mo", "$0.10\u20130.50 per check"],
        ["Card Issuing", "Marqeta / Stripe", "Variable", "Per-card + per-txn"],
        ["Cloud", "AWS / GCP", "$500\u20132,000/mo", "Scale-dependent"],
        ["POS Terminals", "PAX / Sunmi", "$100\u2013300/unit", "One-time"],
        ["SMS", "Africa's Talking", "$50\u2013200/mo", "Per-message"],
        ["Sponsor Bank", "Regional", "Negotiated", "Revenue share"],
    ])

    doc.add_paragraph()
    add_heading_styled(doc, "10. Risk Assessment", level=1)
    add_gold_line(doc)

    styled_table(doc, ["Risk", "Likelihood", "Impact", "Mitigation"], [
        ["BaaS partner delays", "HIGH", "HIGH", "Start Week 1, 2 backup options"],
        ["PCI-DSS timeline", "MEDIUM", "HIGH", "Use BaaS PCI scope"],
        ["Sponsor bank approval", "HIGH", "CRITICAL", "Begin immediately"],
        ["Card network cert", "MEDIUM", "MEDIUM", "Pre-certified BaaS"],
        ["Flutter low-end perf", "LOW", "MEDIUM", "Profile from Week 10"],
        ["Data residency", "MEDIUM", "HIGH", "Multi-region K8s"],
    ])

    doc.add_page_break()

    # ──── SECTION 11: Competitive Advantage ───────────────────────────────────
    add_heading_styled(doc, "11. Our Competitive Advantage", level=1)
    add_gold_line(doc)
    add_body(doc, "What We Bring (Saves 3\u20134 Months vs. Starting from Scratch):", bold=True)

    styled_table(doc, ["Asset", "Value"], [
        ["Apache Fineract Core Banking", "Double-entry GL, loans, savings, multi-currency, multi-tenant"],
        ["9 Payment Providers", "M-Pesa, Airtel, MTN, Flutterwave, Paystack, Cellulant, AT, Razorpay, Stripe"],
        ["React Admin Dashboard", "Full CRUD: clients, loans, savings, reports, mobile money, tax compliance"],
        ["Mobile Banking UX", "50+ proven screens translating directly to Flutter"],
        ["Credit Scoring", "PataScore \u2014 bank statement + M-Pesa analysis"],
    ])

    doc.add_paragraph()
    add_body(doc, "Matching Your Brief Section 4 \u2014 Technical Delivery Capabilities:", bold=True)
    styled_table(doc, ["Your Requirement", "Our Evidence"], [
        ["BaaS Strategy & Discovery", "9 payment providers across Africa + Asia"],
        ["Financial Backend Engineering", "Fineract ledger + 9 payment webhooks + GL posting"],
        ["Secure Mobile Development", "Biometric auth, payment flows, KYC capture"],
        ["Infrastructure Hardening", "Cloud deployment, multi-tenant, API security"],
        ["Fintech Project Orchestration", "Multi-app ecosystem (web + 2 mobile)"],
        ["Financial QA", "Ledger reconciliation, transaction testing"],
    ])

    doc.add_page_break()

    # ──── SECTION 12: Edge Cases ──────────────────────────────────────────────
    add_heading_styled(doc, "12. Edge Case Matrix", level=1)
    add_gold_line(doc)
    add_body(doc, "Representative sample of our Edge Case Matrix for transaction failure states:")

    styled_table(doc, ["Scenario", "Expected Behavior", "Recovery"], [
        ["Card auth timeout", "Hold funds, retry once, release 30min", "Auto-reversal + notify"],
        ["Double STK push", "Idempotency key prevents duplicate", "Return existing txn ID"],
        ["KYC verification fails", "Account in 'Pending' state", "Manual review queue"],
        ["Offline P2P transfer", "Queue locally, sync online", "Conflict detection"],
        ["POS battery dies mid-txn", "Transaction not completed", "'Incomplete' receipt"],
        ["FX rate changes mid-transfer", "Lock rate 60s at initiation", "Re-quote on expiry"],
        ["Sponsor bank downtime", "Degrade to mobile money", "Auto-switch rail"],
        ["Card fraud detected", "Instant freeze + push", "Unfreeze or replace"],
    ])

    doc.add_page_break()

    # ──── SECTION 13: Deliverables & Next Steps ──────────────────────────────
    add_heading_styled(doc, "13. Deliverables & Next Steps", level=1)
    add_gold_line(doc)

    add_heading_styled(doc, "Final Deliverables", level=2)
    add_bullet(doc, " (iOS + Android) \u2014 Full digital banking experience", "Flutter Mobile App")
    add_bullet(doc, " (enhanced) \u2014 Staff operations, merchant management", "React Admin Dashboard")
    add_bullet(doc, " (extended) \u2014 OAuth2, card issuing, KYC, shadow ledger, P2P", "Backend API Layer")
    add_bullet(doc, " \u2014 PCI-compliant cloud, multi-region, CI/CD pipeline", "Infrastructure")
    add_bullet(doc, " \u2014 API docs, architecture diagrams, SOC2, edge case matrix", "Documentation")
    add_bullet(doc, " \u2014 Unit, integration, pen test, stress test reports", "Testing")

    doc.add_paragraph()
    add_heading_styled(doc, "Recommended Next Steps", level=2)
    styled_table(doc, ["#", "Step", "Details"], [
        ["1", "NDA Signing", "Before disclosing brand name and specifications"],
        ["2", "BaaS Discovery Call", "1-week sprint to evaluate 3\u20135 BaaS partners"],
        ["3", "Technical Deep-Dive", "2-hour session with your CTO"],
        ["4", "Contract & Kickoff", "Phase 1 begins immediately after agreement"],
    ])

    doc.add_paragraph()
    doc.add_paragraph()

    # Contact
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Ready to Build the Future of Payments?")
    run.font.size = Pt(16)
    run.font.color.rgb = PRIMARY
    run.bold = True
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Qsoftwares Ltd\ninfo@qsoftwares.org  |  +254 710 401 008\nwww.qsoftwares.org")
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT_BODY
    run.font.name = "Calibri"

    # Save
    doc.save(OUTPUT_PATH)
    size_kb = os.path.getsize(OUTPUT_PATH) / 1024
    print(f"Generated: {OUTPUT_PATH} ({size_kb:.0f} KB)")


if __name__ == "__main__":
    build_docx()
